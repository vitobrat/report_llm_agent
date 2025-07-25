from docx import Document
from io import BytesIO
import logging
import markdown
from bs4 import BeautifulSoup


def convert_md_to_docx(md_content: str) -> BytesIO:
    try:
        # Конвертируем Markdown в HTML
        html_content = markdown.markdown(md_content)

        # Создаем новый DOCX документ
        doc = Document()

        # Парсим HTML и добавляем в документ
        soup = BeautifulSoup(html_content, 'html.parser')

        for element in soup.children:
            if element.name == 'h1':
                doc.add_heading(element.text, level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text, level=2)
            elif element.name == 'h3':
                doc.add_heading(element.text, level=3)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Number')
            elif element.name == 'hr':
                doc.add_paragraph('-' * 50)  # Горизонтальная линия

        # Сохраняем в буфер памяти
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return docx_buffer

    except Exception as e:
        logging.error(f"DOCX conversion error: {e}")
        raise